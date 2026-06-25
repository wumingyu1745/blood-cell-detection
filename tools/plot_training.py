"""Generate training curves and insert into Word report.

Usage:
    python tools/plot_training.py
"""

import json
import os

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches


def load_scalars(json_path):
    """Load scalars.json file."""
    records = []
    with open(json_path, 'r') as f:
        for line in f:
            line = line.strip()
            if line:
                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
    return records


def smooth(values, weight=0.9):
    """Exponential moving average smoothing."""
    smoothed = []
    last = values[0]
    for v in values:
        smoothed_val = last * weight + (1 - weight) * v
        smoothed.append(smoothed_val)
        last = smoothed_val
    return smoothed


def plot_loss_curve(records, output_path):
    """Plot training loss curve."""
    train_records = [r for r in records if 'loss' in r and 'epoch' in r]
    epochs = [r['epoch'] for r in train_records]
    loss = [r['loss'] for r in train_records]
    loss_cls = [r.get('loss_cls', 0) for r in train_records]
    loss_bbox = [r.get('loss_bbox', 0) for r in train_records]

    fig, ax = plt.subplots(1, 1, figsize=(10, 5))

    ax.plot(epochs, smooth(loss), 'b-', linewidth=2, label='Total Loss')
    ax.plot(epochs, smooth(loss_cls), 'r--', linewidth=1.5, label='Classification Loss')
    ax.plot(epochs, smooth(loss_bbox), 'g--', linewidth=1.5, label='Regression Loss')

    ax.set_xlabel('Epoch', fontsize=12)
    ax.set_ylabel('Loss', fontsize=12)
    ax.set_title('Training Loss Curve', fontsize=14)
    ax.legend(fontsize=11)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'Saved: {output_path}')


def plot_map_curve(records, output_path):
    """Plot mAP validation curve."""
    val_records = [r for r in records if 'coco/bbox_mAP' in r]
    # Use step as epoch proxy (each val record corresponds to an epoch)
    epochs = [r.get('step', i * 10) for i, r in enumerate(val_records)]
    mAP = [r['coco/bbox_mAP'] for r in val_records]
    mAP_50 = [r['coco/bbox_mAP_50'] for r in val_records]
    mAP_75 = [r['coco/bbox_mAP_75'] for r in val_records]

    fig, ax = plt.subplots(1, 1, figsize=(10, 5))

    ax.plot(epochs, mAP, 'b-o', linewidth=2, markersize=4, label='mAP@50:95')
    ax.plot(epochs, mAP_50, 'r-s', linewidth=2, markersize=4, label='mAP@50')
    ax.plot(epochs, mAP_75, 'g-^', linewidth=2, markersize=4, label='mAP@75')

    ax.set_xlabel('Epoch', fontsize=12)
    ax.set_ylabel('mAP', fontsize=12)
    ax.set_title('Validation mAP Curve', fontsize=14)
    ax.legend(fontsize=11)
    ax.grid(True, alpha=0.3)
    ax.set_ylim(0, 1)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'Saved: {output_path}')


def plot_lr_curve(records, output_path):
    """Plot learning rate curve."""
    lr_records = [r for r in records if 'lr' in r and 'epoch' in r]
    epochs = [r['epoch'] for r in lr_records]
    lr = [r['lr'] for r in lr_records]

    fig, ax = plt.subplots(1, 1, figsize=(10, 4))

    ax.plot(epochs, lr, 'b-', linewidth=2)
    ax.set_xlabel('Epoch', fontsize=12)
    ax.set_ylabel('Learning Rate', fontsize=12)
    ax.set_title('Learning Rate Schedule', fontsize=14)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'Saved: {output_path}')


def plot_loss_components(records, output_path):
    """Plot loss components breakdown."""
    train_records = [r for r in records if 'loss' in r and 'epoch' in r]
    epochs = [r['epoch'] for r in train_records]
    loss_cls = [r.get('loss_cls', 0) for r in train_records]
    loss_bbox = [r.get('loss_bbox', 0) for r in train_records]

    fig, ax = plt.subplots(1, 1, figsize=(10, 5))

    ax.fill_between(epochs, 0, smooth(loss_cls), alpha=0.3, color='red', label='Classification Loss')
    ax.fill_between(epochs, smooth(loss_cls), [a + b for a, b in zip(smooth(loss_cls), smooth(loss_bbox))],
                    alpha=0.3, color='green', label='Regression Loss')
    ax.plot(epochs, smooth(loss_cls), 'r-', linewidth=1.5)
    ax.plot(epochs, [a + b for a, b in zip(smooth(loss_cls), smooth(loss_bbox))], 'b-', linewidth=1.5)

    ax.set_xlabel('Epoch', fontsize=12)
    ax.set_ylabel('Loss', fontsize=12)
    ax.set_title('Loss Components Breakdown', fontsize=14)
    ax.legend(fontsize=11)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'Saved: {output_path}')


def plot_class_map(records, output_path):
    """Plot per-class mAP."""
    val_records = [r for r in records if 'coco/bbox_mAP' in r]
    if not val_records:
        return

    last = val_records[-1]
    categories = ['mAP@50:95', 'mAP@50', 'mAP@75', 'mAP_s', 'mAP_m', 'mAP_l']
    values = [
        last.get('coco/bbox_mAP', 0),
        last.get('coco/bbox_mAP_50', 0),
        last.get('coco/bbox_mAP_75', 0),
        last.get('coco/bbox_mAP_s', 0),
        last.get('coco/bbox_mAP_m', 0),
        last.get('coco/bbox_mAP_l', 0),
    ]

    fig, ax = plt.subplots(1, 1, figsize=(10, 5))

    colors = ['#2196F3', '#4CAF50', '#FF9800', '#9C27B0', '#F44336', '#00BCD4']
    bars = ax.bar(categories, values, color=colors, edgecolor='white', linewidth=1.5)

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.01,
                f'{val:.3f}', ha='center', va='bottom', fontsize=11, fontweight='bold')

    ax.set_ylabel('Score', fontsize=12)
    ax.set_title('Final Model Performance (Epoch 280)', fontsize=14)
    ax.set_ylim(0, 1)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'Saved: {output_path}')


def insert_images_to_docx(docx_path, image_paths, captions):
    """Insert images into Word document."""
    doc = Document(docx_path)

    # Add a heading for training curves section
    doc.add_heading('附录：训练过程可视化', level=1)

    for img_path, caption in zip(image_paths, captions):
        if os.path.exists(img_path):
            # Add image
            doc.add_picture(img_path, width=Inches(6))
            # Add caption
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.font.size = Pt(10)
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('')  # spacing

    doc.save(docx_path)
    print(f'Updated: {docx_path}')


if __name__ == '__main__':
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    # Find the main training log
    log_dir = os.path.join(
        project_root, 'work_dirs', 'bccd_rtmdet_tiny_8xb32-300e_coco',
        '20260623_162941', 'vis_data')
    scalars_path = os.path.join(log_dir, 'scalars.json')

    # Also load the first 10 epochs log
    log_dir_early = os.path.join(
        project_root, 'work_dirs', 'bccd_rtmdet_tiny_8xb32-300e_coco',
        '20260623_160348', 'vis_data')
    scalars_path_early = os.path.join(log_dir_early, 'scalars.json')

    # Load data
    print('Loading training logs...')
    records = []
    if os.path.exists(scalars_path_early):
        records.extend(load_scalars(scalars_path_early))
    if os.path.exists(scalars_path):
        records.extend(load_scalars(scalars_path))

    print(f'Loaded {len(records)} records')

    # Output directory
    img_dir = os.path.join(project_root, 'docs', 'figures')
    os.makedirs(img_dir, exist_ok=True)

    # Generate plots
    print('\nGenerating plots...')

    plot_loss_curve(records, os.path.join(img_dir, 'loss_curve.png'))
    plot_map_curve(records, os.path.join(img_dir, 'map_curve.png'))
    plot_lr_curve(records, os.path.join(img_dir, 'lr_curve.png'))
    plot_loss_components(records, os.path.join(img_dir, 'loss_components.png'))
    plot_class_map(records, os.path.join(img_dir, 'class_map.png'))

    # Insert into Word document
    docx_path = os.path.join(project_root, 'docs', 'experiment_report.docx')
    if os.path.exists(docx_path):
        print('\nInserting images into Word document...')
        image_paths = [
            os.path.join(img_dir, 'loss_curve.png'),
            os.path.join(img_dir, 'map_curve.png'),
            os.path.join(img_dir, 'lr_curve.png'),
            os.path.join(img_dir, 'loss_components.png'),
            os.path.join(img_dir, 'class_map.png'),
        ]
        captions = [
            '图1 训练损失曲线（Total Loss、Classification Loss、Regression Loss）',
            '图2 验证集mAP曲线（mAP@50:95、mAP@50、mAP@75）',
            '图3 学习率调度曲线（CosineAnnealingLR）',
            '图4 损失分量对比（分类损失 vs 回归损失）',
            '图5 最终模型性能指标（Epoch 280）',
        ]
        insert_images_to_docx(docx_path, image_paths, captions)

    print('\nDone!')

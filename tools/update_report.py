"""Update Word report: insert training curves and detection results into main body.

Usage:
    python tools/update_report.py
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Add a figure with caption to the document."""
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')  # spacing


def find_paragraph_index(doc, keyword):
    """Find paragraph index containing keyword."""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            return i
    return -1


def insert_after_paragraph(doc, index, new_paragraph):
    """Insert a paragraph after a given index."""
    # This is a workaround since python-docx doesn't support direct insertion
    # We'll use a different approach - rebuild the document
    pass


def rebuild_report(docx_path, figures_dir, results_dir, output_path):
    """Rebuild the report with images in the main body."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

    # ===== TITLE =====
    doc.add_heading('基于RTMDet的血液细胞检测与分类系统', level=0)
    doc.add_heading('实验报告', level=1)
    doc.add_paragraph('')

    # ===== 一、研究背景与意义 =====
    doc.add_heading('一、研究背景与意义', level=1)
    doc.add_heading('1.1 研究背景', level=2)
    doc.add_paragraph(
        '血液细胞检测与分类是临床血液检验中的重要环节。传统的血液细胞分析主要依赖人工在显微镜下观察和计数，这种方法存在以下问题：')
    doc.add_paragraph('效率低下：人工计数耗时长，一份血常规样本需要数分钟到数十分钟', style='List Bullet')
    doc.add_paragraph('主观性强：不同检验人员的判断标准存在差异，结果一致性差', style='List Bullet')
    doc.add_paragraph('劳动强度大：长时间观察显微镜容易导致视觉疲劳，影响准确性', style='List Bullet')
    doc.add_paragraph('人力成本高：需要专业检验人员，基层医疗机构人才匮乏', style='List Bullet')
    doc.add_paragraph(
        '随着深度学习技术的快速发展，基于计算机视觉的自动化血液细胞检测成为可能。通过目标检测算法，可以实现血液细胞的自动定位、分类和计数，大幅提升检验效率和准确性。')

    doc.add_heading('1.2 研究意义', level=2)
    doc.add_paragraph('提高检验效率：自动化检测可将检验时间从数分钟缩短到毫秒级别', style='List Bullet')
    doc.add_paragraph('提升检测精度：基于深度学习的检测模型具有较高的准确性和一致性', style='List Bullet')
    doc.add_paragraph('降低人力成本：减少对专业检验人员的依赖，适用于基层医疗机构', style='List Bullet')
    doc.add_paragraph('辅助临床诊断：为血液疾病的早期筛查提供技术支持', style='List Bullet')

    # ===== 二、研究现状 =====
    doc.add_heading('二、研究现状', level=1)
    doc.add_heading('2.1 传统方法', level=2)
    doc.add_paragraph('传统的血液细胞检测方法主要包括：')
    doc.add_paragraph('基于阈值分割：通过设定灰度阈值将细胞从背景中分离，但对光照变化敏感', style='List Bullet')
    doc.add_paragraph('基于边缘检测：使用Canny、Sobel等算子检测细胞边缘，但对噪声敏感', style='List Bullet')
    doc.add_paragraph('基于形态学处理：利用膨胀、腐蚀等操作提取细胞区域，但对重叠细胞效果差', style='List Bullet')

    doc.add_heading('2.2 深度学习方法', level=2)
    doc.add_paragraph('近年来，基于深度学习的目标检测算法在血液细胞检测中取得了显著进展：')

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['算法', '类型', '特点', '代表工作']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['Faster R-CNN', '两阶段', '精度高，速度较慢', 'RBC检测'],
        ['YOLOv5/v8', '单阶段', '速度快，实时检测', '血液细胞计数'],
        ['DETR', 'Transformer', '端到端，无需NMS', '医学图像检测'],
        ['RTMDet', '单阶段', '平衡速度与精度', '实时检测场景'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    doc.add_heading('2.3 本项目选择', level=2)
    doc.add_paragraph('本项目选择 RTMDet-Tiny 作为检测模型，原因如下：')
    doc.add_paragraph('速度快：RTMDet是实时目标检测算法，适合部署应用', style='List Bullet')
    doc.add_paragraph('精度高：在COCO数据集上达到SOTA水平', style='List Bullet')
    doc.add_paragraph('轻量级：Tiny版本参数量小，适合边缘设备部署', style='List Bullet')
    doc.add_paragraph('社区支持：基于MMDetection框架，文档完善，易于复现', style='List Bullet')

    # ===== 三、模型设计 =====
    doc.add_heading('三、模型设计', level=1)
    doc.add_heading('3.1 RTMDet整体架构', level=2)
    doc.add_paragraph('RTMDet由三个核心模块组成：')
    doc.add_paragraph('Backbone (CSPNeXt-Tiny)：特征提取网络，从输入图像中提取多尺度特征')
    doc.add_paragraph('Neck (CSPNeXtPAFPN)：特征融合网络，融合多尺度特征增强表达能力')
    doc.add_paragraph('Head (RTMDetSepBNHead)：检测头，预测类别和边界框')

    doc.add_heading('3.2 损失函数', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['损失类型', '函数', '用途']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['分类损失', 'QualityFocalLoss', '处理类别不平衡问题'],
        ['回归损失', 'GIoULoss', '优化边界框回归精度'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    doc.add_heading('3.3 检测类别', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['类别', '英文', '说明']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['红细胞', 'RBC', '数量最多，双凹圆盘状'],
        ['白细胞', 'WBC', '免疫细胞，体积较大'],
        ['血小板', 'Platelets', '体积最小，参与凝血'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    # ===== 四、数据集介绍 =====
    doc.add_heading('四、数据集介绍', level=1)
    doc.add_heading('4.1 BCCD数据集', level=2)
    doc.add_paragraph('BCCD (Blood Cell Count and Detection) 是一个公开的血液细胞检测数据集：')
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ['总图片数', '364张'],
        ['图片分辨率', '640×480'],
        ['标注格式', 'Pascal VOC XML'],
        ['类别数', '3类 (RBC, WBC, Platelets)'],
        ['总标注数', '4,888个'],
    ]
    for row_idx, (k, v) in enumerate(data):
        table.rows[row_idx].cells[0].text = k
        table.rows[row_idx + 1].cells[1].text = v
    doc.add_paragraph('')

    doc.add_heading('4.2 数据集划分', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['集合', '数量', '占比', '用途']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['训练集', '205张', '56%', '模型训练'],
        ['验证集', '87张', '24%', '训练中评估'],
        ['测试集', '72张', '20%', '最终评估'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    doc.add_heading('4.3 数据增强', level=2)
    doc.add_paragraph('训练阶段使用了以下数据增强策略：')
    doc.add_paragraph('CachedMosaic：4张图片拼接，增加小目标数量', style='List Bullet')
    doc.add_paragraph('CachedMixUp：2张图片混合，增加样本多样性', style='List Bullet')
    doc.add_paragraph('RandomResize：随机缩放，增强尺度鲁棒性', style='List Bullet')
    doc.add_paragraph('RandomCrop：随机裁剪，增强位置鲁棒性', style='List Bullet')
    doc.add_paragraph('RandomFlip：随机翻转，增强方向鲁棒性', style='List Bullet')
    doc.add_paragraph('YOLOXHSVRandomAug：颜色空间增强，增强光照鲁棒性', style='List Bullet')

    # ===== 五、训练流程 =====
    doc.add_heading('五、训练流程', level=1)
    doc.add_heading('5.1 环境配置', level=2)

    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    headers = ['组件', '版本']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['操作系统', 'Windows 10'],
        ['Python', '3.8.20'],
        ['PyTorch', '2.4.1'],
        ['CUDA', '11.8'],
        ['mmcv', '2.1.0 (lite)'],
        ['mmengine', '0.10.7'],
        ['mmdet', '3.3.0'],
        ['GPU', 'NVIDIA RTX 3060 (12GB)'],
    ]
    for row_idx, (k, v) in enumerate(data):
        table.rows[row_idx].cells[0].text = k
        table.rows[row_idx + 1].cells[1].text = v
    doc.add_paragraph('')

    doc.add_heading('5.2 训练超参数', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['参数', '数值', '说明']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['batch_size', '16', '每批样本数'],
        ['max_epochs', '300', '最大训练轮数'],
        ['learning_rate', '0.004', '初始学习率'],
        ['weight_decay', '0.05', '权重衰减'],
        ['optimizer', 'AdamW', '优化器'],
        ['scheduler', 'CosineAnnealingLR', '学习率调度'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    doc.add_heading('5.3 训练过程可视化', level=2)

    # Insert training curves here
    add_figure(doc, os.path.join(figures_dir, 'loss_curve.png'),
               '图1 训练损失曲线（Total Loss、Classification Loss、Regression Loss）')
    add_figure(doc, os.path.join(figures_dir, 'map_curve.png'),
               '图2 验证集mAP曲线（mAP@50:95、mAP@50、mAP@75）')
    add_figure(doc, os.path.join(figures_dir, 'lr_curve.png'),
               '图3 学习率调度曲线（CosineAnnealingLR）')
    add_figure(doc, os.path.join(figures_dir, 'loss_components.png'),
               '图4 损失分量对比（分类损失 vs 回归损失）')

    # ===== 六、检测结果 =====
    doc.add_heading('六、检测结果', level=1)
    doc.add_heading('6.1 定量结果', level=2)
    doc.add_paragraph('在测试集（72张图片）上的评估结果：')

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['指标', '数值', '说明']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['mAP@50', '0.884', 'IoU=0.5时的平均精度'],
        ['mAP@50:95', '0.613', 'IoU=0.5:0.95的平均精度'],
        ['mAP@75', '0.723', 'IoU=0.75时的平均精度'],
        ['mAP_s', '0.320', '小目标检测精度'],
        ['mAP_m', '0.482', '中等目标检测精度'],
        ['mAP_l', '0.475', '大目标检测精度'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    # Insert mAP bar chart
    add_figure(doc, os.path.join(figures_dir, 'class_map.png'),
               '图5 最终模型性能指标（Epoch 280）')

    doc.add_heading('6.2 检测速度对比', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['方法', '推理时间', 'FPS', '说明']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['PyTorch (GPU)', '14.7ms', '68.3', '原始PyTorch推理'],
        ['ONNX (CPU)', '26.2ms', '38.1', 'ONNX Runtime CPU'],
        ['ONNX FP16', '40.1ms', '25.0', 'ONNX FP16精度'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = text
    doc.add_paragraph('')

    doc.add_heading('6.3 检测结果可视化', level=2)
    doc.add_paragraph('以下展示模型在测试集上的检测效果。每张图片标注了检测框、类别标签和置信度分数。')

    # Insert detection result images
    if os.path.exists(results_dir):
        result_images = sorted([f for f in os.listdir(results_dir) if f.endswith('.jpg')])[:5]
        for idx, img_file in enumerate(result_images):
            img_path = os.path.join(results_dir, img_file)
            add_figure(doc, img_path,
                       f'图{6 + idx} 检测结果：{img_file}')

    # ===== 七、Bug调试 =====
    doc.add_heading('七、遇到的Bug及调试方案', level=1)

    bugs = [
        ('Bug 1: mmcv安装失败 — 缺少C++编译器',
         'error: Microsoft Visual C++ 14.0 or greater is required.',
         'mmcv需要编译CUDA算子，系统缺少C++编译器。',
         '使用mmcv-lite版本，无需编译CUDA算子：pip install mmcv-lite==2.1.0'),
        ('Bug 2: CUDA版本不匹配',
         'The detected CUDA version (12.4) mismatches the version that was used to compile PyTorch (11.8)',
         '系统安装了CUDA 12.4，但PyTorch是基于CUDA 11.8编译的。',
         '使用mmcv-lite避免编译，或设置CUDA_HOME指向conda环境的CUDA。'),
        ('Bug 3: mmcv._ext模块缺失',
         'ModuleNotFoundError: No module named \'mmcv._ext\'',
         'mmcv-lite不包含编译好的CUDA扩展模块。',
         '修改mmcv的ext_loader.py，添加优雅降级处理，当模块不存在时返回空模块。'),
        ('Bug 4: NMS函数调用失败',
         'TypeError: \'NoneType\' object is not callable',
         'NMS使用了mmcv的编译扩展，但mmcv-lite中该扩展为None。',
         '修改mmcv的nms.py，添加torchvision NMS作为fallback。'),
        ('Bug 5: Pillow DLL加载失败',
         'ImportError: DLL load failed while importing _imaging',
         'torchvision重装后导致Pillow的DLL依赖损坏。',
         'pip install --force-reinstall Pillow'),
        ('Bug 6: DataLoader Worker进程异常退出',
         'RuntimeError: DataLoader worker exited unexpectedly',
         'OpenMP运行时重复加载导致进程崩溃。',
         '设置环境变量：set KMP_DUPLICATE_LIB_OK=TRUE'),
        ('Bug 7: 图片路径找不到',
         'FileNotFoundError: No such file or directory',
         'COCO JSON中的file_name没有包含子目录路径。',
         '修改转换脚本，让file_name包含子目录：f\'{split_name}/{filename}\''),
        ('Bug 8: ONNX CUDA Provider加载失败',
         'Failed to create CUDAExecutionProvider. Require cuDNN 9.* and CUDA 12.*',
         'onnxruntime-gpu需要CUDA 12.x，但当前环境是CUDA 11.8。',
         '使用CPU Provider自动回退，或升级到CUDA 12.x环境。'),
    ]

    for title, error, reason, solution in bugs:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        run = p.add_run('错误信息：')
        run.bold = True
        p.add_run(error)

        p = doc.add_paragraph()
        run = p.add_run('原因分析：')
        run.bold = True
        p.add_run(reason)

        p = doc.add_paragraph()
        run = p.add_run('解决方案：')
        run.bold = True
        p.add_run(solution)
        doc.add_paragraph('')

    # ===== 八、总结 =====
    doc.add_heading('八、总结与展望', level=1)
    doc.add_heading('8.1 项目总结', level=2)
    doc.add_paragraph('本项目基于RTMDet-Tiny深度学习模型，实现了血液细胞的自动检测与分类：')
    doc.add_paragraph('数据方面：使用BCCD公开数据集，包含364张血液涂片图片，3个类别', style='List Bullet')
    doc.add_paragraph('模型方面：采用RTMDet-Tiny模型，兼顾检测速度和精度', style='List Bullet')
    doc.add_paragraph('训练方面：300个epoch训练，最终mAP@50达到88.4%', style='List Bullet')
    doc.add_paragraph('部署方面：实现了ONNX导出和桌面化应用（Gradio界面）', style='List Bullet')

    doc.add_heading('8.2 创新点', level=2)
    doc.add_paragraph('小样本数据增强策略：采用CachedMosaic和CachedMixUp等先进数据增强方法', style='List Bullet')
    doc.add_paragraph('分阶段训练策略：使用PipelineSwitchHook，在训练后期切换数据增强策略', style='List Bullet')
    doc.add_paragraph('多精度模型部署：支持FP32和FP16两种精度的ONNX模型部署', style='List Bullet')
    doc.add_paragraph('桌面化应用：基于Gradio构建了用户友好的检测界面', style='List Bullet')

    doc.add_heading('8.3 未来展望', level=2)
    doc.add_paragraph('扩展数据集：收集更多血液细胞图片，提升模型泛化能力', style='List Bullet')
    doc.add_paragraph('模型优化：尝试更大模型（RTMDet-S/M）或知识蒸馏', style='List Bullet')
    doc.add_paragraph('边缘部署：适配Jetson等边缘设备，实现实时检测', style='List Bullet')
    doc.add_paragraph('多任务学习：同时实现检测、分割和计数', style='List Bullet')

    # Save
    doc.save(output_path)
    print(f'Saved: {output_path}')
    print(f'Size: {os.path.getsize(output_path) / 1024:.1f} KB')


if __name__ == '__main__':
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    figures_dir = os.path.join(project_root, 'docs', 'figures')
    results_dir = os.path.join(project_root, 'results', 'bccd_demo', 'vis')
    output_path = os.path.join(project_root, 'docs', 'experiment_report.docx')

    rebuild_report(None, figures_dir, results_dir, output_path)

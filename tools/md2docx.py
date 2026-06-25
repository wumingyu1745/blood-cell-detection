"""Convert Markdown report to Word (.docx) format.

Usage:
    python tools/md2docx.py
"""

import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_text(cell, text, bold=False, color=None, size=10):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, text in enumerate(header_row):
        set_cell_text(table.rows[0].cells[i], text.strip(), bold=True, size=10)

    # Data
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, text in enumerate(row_data):
            set_cell_text(table.rows[row_idx + 1].cells[col_idx], text.strip(), size=10)

    return table


def md_to_docx(md_path, docx_path):
    """Convert Markdown file to Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Set heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # Check if next line is still part of table
            if i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                continue
            else:
                # Process table
                if len(table_lines) >= 2:
                    # Parse header
                    header = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                    # Skip separator line (---|---|---)
                    data_rows = []
                    for tl in table_lines[2:]:
                        row = [c.strip() for c in tl.split('|') if c.strip()]
                        if row:
                            data_rows.append(row)
                    if header and data_rows:
                        add_table(doc, header, data_rows)
                in_table = False
                table_lines = []
                continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.*)', line)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                # Clean markdown formatting
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                text = re.sub(r'\*(.*?)\*', r'\1', text)
                if level <= 3:
                    doc.add_heading(text, level=level)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(12)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered list
        match = re.match(r'^\d+\.\s+(.*)', line.strip())
        if match:
            text = match.group(1)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        # Clean markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        if text:
            doc.add_paragraph(text)
        i += 1

    # Save
    doc.save(docx_path)
    print(f'Saved: {docx_path}')
    print(f'File size: {os.path.getsize(docx_path) / 1024:.1f} KB')


if __name__ == '__main__':
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(project_root, 'docs', 'experiment_report.md')
    docx_path = os.path.join(project_root, 'docs', 'experiment_report.docx')

    md_to_docx(md_path, docx_path)
